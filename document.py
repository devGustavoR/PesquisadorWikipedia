import wikipedia
import re
from docx import Document
from PySimpleGUI import PySimpleGUI as sg

sg.theme('Dark')
layout = [
    [sg.Text('Nome'), sg.Input(key='nome')],
    [sg.Text('Tema'), sg.Input(key='tema')],
    [sg.Button('Pesquisar')]
]

janela = sg.Window('Pesquisador Wikipedia', layout)
wikipedia.set_lang('pt')

while True:
    eventos, valores = janela.read()
    if eventos == sg.WINDOW_CLOSED:
        break
    if eventos == 'Pesquisar':
        wiki = wikipedia.page(valores['tema'])


        text = wiki.content
        text = re.sub(r'==','',text)
        text = re.sub(r'=','',text)
        split = text.split('Veja também',1)
        text = split[0]

        print(text)


        document = Document()
        paragraph = document.add_heading(valores['tema'],0)
        paragraph.alignment = 1

        paragraph = document.add_paragraph('   '+text)
        paragraph.alignment = 2
        document.save(valores['tema'] + '.docx')

        break

    else:
        print("Nome do projeto inválido")
        title = input("Digite outro nome:")

